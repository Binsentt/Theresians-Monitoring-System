"""Create the two reviewed Grade 1 / Easy / Basic Addition teacher documents.

These documents intentionally use the same simple numbered-question format the
Fixed Questions DOCX/PDF parser accepts. They are upload candidates only; this
script never calls the application API or writes to production.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIRECTORY = Path(__file__).parent
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)


QUESTION_SETS = {
    "grade1-easy-basic-addition-set-a.docx": [
        ("What is 2 + 3?", ["4", "5", "6", "7"], "B"),
        ("What is 1 + 4?", ["3", "4", "5", "6"], "C"),
        ("What is 5 + 2?", ["6", "7", "8", "9"], "B"),
        ("What is 3 + 3?", ["5", "6", "7", "8"], "B"),
        ("What is 4 + 4?", ["6", "7", "8", "9"], "C"),
    ],
    "grade1-easy-basic-addition-set-b.docx": [
        ("What is 1 + 2?", ["2", "3", "4", "5"], "B"),
        ("What is 2 + 4?", ["5", "6", "7", "8"], "B"),
        ("What is 3 + 4?", ["5", "6", "7", "8"], "C"),
        ("What is 5 + 3?", ["6", "7", "8", "9"], "C"),
        ("What is 6 + 2?", ["7", "8", "9", "10"], "B"),
    ],
}


def set_run_font(run, name="Calibri", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_text(document, text, *, size=11, bold=False, color=None, after=6):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def configure_document(document):
    section = document.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25


def build_document(filename, questions):
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Theresian's Quest — Fixed Questions")
    set_run_font(title_run, size=18, bold=True, color=(46, 116, 181))

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Grade 1 · Easy · Basic Addition · 5 Questions")
    set_run_font(subtitle_run, size=11, bold=False, color=(31, 77, 120))

    add_text(
        document,
        "Teacher upload instructions: select Fixed Questions, choose Grade 1, "
        "Easy, and Basic Addition. Keep every question, four choices, and Answer line unchanged.",
        size=10,
        color=(58, 58, 58),
        after=12,
    )

    for index, (question, choices, answer) in enumerate(questions, start=1):
        question_paragraph = add_text(document, f"{index}. {question}", size=11, bold=True, after=2)
        question_paragraph.paragraph_format.keep_with_next = True
        for letter, choice in zip("ABCD", choices):
            option_paragraph = add_text(document, f"{letter}. {choice}", size=11, after=1)
            option_paragraph.paragraph_format.left_indent = Inches(0.25)
        answer_paragraph = add_text(document, f"Answer: {answer}", size=11, bold=True, after=9)
        answer_paragraph.paragraph_format.left_indent = Inches(0.25)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Theresian's Quest — Teacher Fixed Questions")
    set_run_font(footer_run, size=9, color=(89, 89, 89))

    document.save(OUTPUT_DIRECTORY / filename)


def main():
    for filename, questions in QUESTION_SETS.items():
        build_document(filename, questions)


if __name__ == "__main__":
    main()
